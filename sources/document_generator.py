#!/bin/python3

from os import getenv

from docx.table import Table

from sources.fetch_github import Issue, get_issues
from sources.json_file import get_json_from_file
from sources.document_utils import (add_title_center, add_picture_center, add_title,
                                    save, add_paragraph_indent, page_break, create_array, blank_line, add_paragraph, create_header)
from datetime import date

CONFIG_PATH: str = getenv("CONFIG_PATH")
RESOURCES_FOLDER: str = getenv("RESOURCES_FOLDER")


def create_cover_page(config: dict) -> None:
    cover_image: str = f'{RESOURCES_FOLDER}/{config["cover-image"]["path"]}'
    size: dict[str, float] = config["cover-image"]["size"]
    add_picture_center(cover_image, width=size["width"], height=size["height"])
    page_break()


def create_document_description(description: dict) -> None:
    create_header()
    add_title("Description du document")
    array: Table = create_array(9, 2)
    array.style = 'Table Grid'
    array.rows[0].cells[0].text = "Titre"
    array.rows[0].cells[1].text = description["Titre"]
    array.rows[1].cells[0].text = "Objet"
    array.rows[1].cells[1].text = description["Objet"]
    array.rows[2].cells[0].text = "Auteur"
    array.rows[2].cells[1].text = description["Auteur"]
    array.rows[3].cells[0].text = "Responsable"
    array.rows[3].cells[1].text = description["Responsable"]
    array.rows[4].cells[0].text = "E-mail"
    array.rows[4].cells[1].text = description["E-mail"]
    array.rows[5].cells[0].text = "Mots-clés"
    array.rows[5].cells[1].text = ", ".join(list(description["Mots-clés"]))
    array.rows[6].cells[0].text = "Promotion"
    array.rows[6].cells[1].text = description["Promotion"]
    array.rows[7].cells[0].text = "Date de mise à jour"
    array.rows[7].cells[1].text = date.today().strftime("%d/%m/%Y")
    array.rows[8].cells[0].text = "Version du modèle"
    array.rows[8].cells[1].text = "1.0"
    page_break()


def create_revision_table() -> None:
    add_title("Tableau de révision")
    array: Table = create_array(2, 5)
    array.style = 'Table Grid'
    header_values: list[str] = ["Date", "Version", "Auteur", "Section(s)", "Commentaires"]
    for index in range(len(array.rows[0].cells)):
        array.rows[0].cells[index].text = header_values[index]
    array.rows[1].cells[0].text = date.today().strftime("%d/%m/%Y")
    array.rows[1].cells[1].text = "1.0"
    array.rows[1].cells[2].text = "WhaleMusic"
    array.rows[1].cells[3].text = "Toutes"
    array.rows[1].cells[4].text = "Première version"
    page_break()


def create_organigramme_livrable(image_infos: dict) -> None:
    add_title("Organigramme des livrables")
    image_path: str = f'{RESOURCES_FOLDER}/{image_infos["path"]}'
    size: dict[str, float] = image_infos["size"]
    add_picture_center(image_path, size["width"], size["height"])
    page_break()


def create_livrable_map(livrable_maps: list[dict]) -> None:
    add_title("Carte des livrables")
    for livrable in livrable_maps:
        add_title(livrable["name"], 2)
        image_path: str = f'{RESOURCES_FOLDER}/{livrable["path"]}'
        size: dict[str, float] = livrable["size"]
        add_picture_center(image_path, size["width"], size["height"])


def create_card(issue: Issue) -> None:
    array: Table = create_array(7, 2)
    content: dict = issue.body
    array.style = 'Table Grid'
    array.rows[0].cells[0].text = issue.title
    array.rows[0].cells[0].merge(array.rows[0].cells[1])
    array.rows[1].cells[0].text = "En tant que:"
    array.rows[2].cells[0].text = content["En tant que:"]
    array.rows[1].cells[1].text = "Je veux:"
    array.rows[2].cells[1].text = content["Je veux:"]
    array.rows[3].cells[0].text = f'Description: {content["Description:"]}'
    array.rows[3].cells[0].merge(array.rows[3].cells[1])
    array.rows[4].cells[0].text = "Definition of done:" "\n- " + "\n- ".join(content["Definition of done:"])
    array.rows[4].cells[0].merge(array.rows[4].cells[1])
    array.rows[5].cells[0].text = "Charge estimée"
    array.rows[5].cells[1].text = content["Charge estimée"]
    array.rows[6].cells[0].text = "Assignés:"
    array.rows[6].cells[1].text = ", ".join(issue.assignees)



def create_document():
    config = get_json_from_file(CONFIG_PATH)
    create_cover_page(config)
    create_document_description(config["description"])
    create_revision_table()
    create_organigramme_livrable(config["organigramme-livrable"])
    create_livrable_map(config["livrable-maps"])
    save(RESOURCES_FOLDER + "/" + config["document_name"])
